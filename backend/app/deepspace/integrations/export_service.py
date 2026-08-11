import io

from bs4 import BeautifulSoup
from bs4.element import Tag
from docx import Document
from fpdf import FPDF

_PDF_TEXT_REPLACEMENTS = {
    "\u00a0": " ",
    "\u2010": "-",
    "\u2011": "-",
    "\u2012": "-",
    "\u2013": "-",
    "\u2014": "-",
    "\u2018": "'",
    "\u2019": "'",
    "\u201c": '"',
    "\u201d": '"',
    "\u2022": "-",
    "\u2026": "...",
}


def _safe_pdf_text(value: object) -> str:
    """Normalize text for FPDF core fonts, which only support latin-1."""
    text = str(value or "")
    for source, target in _PDF_TEXT_REPLACEMENTS.items():
        text = text.replace(source, target)
    return text.encode("latin-1", errors="replace").decode("latin-1")


class _PremiumPDF(FPDF):  # type: ignore
    """Custom FPDF subclass with premium styling for DeepSpace notes."""

    def __init__(self, title: str = "DeepSpace Note") -> None:
        super().__init__()
        self._doc_title = _safe_pdf_text(title)
        self.set_auto_page_break(auto=True, margin=25)

    def header(self) -> None:
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(160, 160, 160)
        self.cell(0, 8, self._doc_title, align="R")
        self.ln(4)
        self.set_draw_color(220, 220, 220)
        self.line(10, self.get_y(), self.w - 10, self.get_y())
        self.ln(6)

    def footer(self) -> None:
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(160, 160, 160)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


class DeepSpaceExportService:
    """Generates premium PDF and DOCX exports from DeepSpace note HTML."""

    # ------------------------------------------------------------------ PDF
    def generate_pdf(self, html_content: str, title: str = "DeepSpace Note") -> io.BytesIO:
        """Build a clean, professionally styled PDF from the note HTML."""
        pdf = _PremiumPDF(title=title)
        pdf.alias_nb_pages()
        pdf.add_page()
        pdf.set_margins(20, 20, 20)

        # Title page header
        pdf.set_font("Helvetica", "B", 22)
        pdf.set_text_color(30, 30, 30)
        pdf.cell(0, 14, _safe_pdf_text(title), new_x="LMARGIN", new_y="NEXT", align="C")
        pdf.ln(2)
        pdf.set_draw_color(15, 95, 121)  # brand cyan
        pdf.set_line_width(0.6)
        pdf.line(60, pdf.get_y(), pdf.w - 60, pdf.get_y())
        pdf.ln(12)

        # Parse HTML and render blocks
        soup = BeautifulSoup(html_content, "html.parser")
        self._render_elements(pdf, soup)

        output = io.BytesIO()
        pdf.output(output)
        output.seek(0)
        return output

    def _render_elements(self, pdf: _PremiumPDF, soup: BeautifulSoup | Tag) -> None:
        """Walk top-level elements and render them into the PDF."""
        for el in soup.children:
            if not isinstance(el, Tag):
                # Raw text node
                text = str(el).strip()
                if text:
                    pdf.set_font("Helvetica", "", 11)
                    pdf.set_text_color(51, 51, 51)
                    pdf.multi_cell(0, 6, _safe_pdf_text(text))
                    pdf.ln(3)
                continue

            tag = el.name.lower()

            if tag == "h1":
                pdf.set_font("Helvetica", "B", 20)
                pdf.set_text_color(26, 26, 26)
                pdf.multi_cell(0, 10, _safe_pdf_text(el.get_text(strip=True)))
                pdf.ln(2)
                pdf.set_draw_color(230, 230, 230)
                pdf.line(20, pdf.get_y(), pdf.w - 20, pdf.get_y())
                pdf.ln(6)

            elif tag == "h2":
                pdf.set_font("Helvetica", "B", 16)
                pdf.set_text_color(44, 62, 80)
                pdf.multi_cell(0, 8, _safe_pdf_text(el.get_text(strip=True)))
                pdf.ln(4)

            elif tag == "h3":
                pdf.set_font("Helvetica", "B", 13)
                pdf.set_text_color(52, 73, 94)
                pdf.multi_cell(0, 7, _safe_pdf_text(el.get_text(strip=True)))
                pdf.ln(3)

            elif tag in ("h4", "h5", "h6"):
                pdf.set_font("Helvetica", "B", 11)
                pdf.set_text_color(52, 73, 94)
                pdf.multi_cell(0, 6, _safe_pdf_text(el.get_text(strip=True)))
                pdf.ln(3)

            elif tag == "p":
                pdf.set_font("Helvetica", "", 11)
                pdf.set_text_color(51, 51, 51)
                pdf.multi_cell(0, 6, _safe_pdf_text(el.get_text(strip=True)))
                pdf.ln(4)

            elif tag in ("ul", "ol"):
                items = el.find_all("li", recursive=False)
                for idx, li in enumerate(items, 1):
                    bullet = f"  {idx}." if tag == "ol" else "  -"
                    pdf.set_font("Helvetica", "", 11)
                    pdf.set_text_color(51, 51, 51)
                    pdf.cell(10, 6, bullet)
                    pdf.multi_cell(0, 6, _safe_pdf_text(li.get_text(strip=True)))
                    pdf.ln(2)
                pdf.ln(3)

            elif tag in ("pre", "code"):
                pdf.set_font("Courier", "", 9)
                pdf.set_text_color(80, 80, 80)
                pdf.set_fill_color(248, 249, 250)
                text = el.get_text()
                pdf.multi_cell(0, 5, _safe_pdf_text(text), fill=True)
                pdf.ln(4)

            elif tag == "blockquote":
                pdf.set_font("Helvetica", "I", 11)
                pdf.set_text_color(106, 115, 125)
                pdf.set_draw_color(15, 95, 121)
                x = pdf.get_x()
                y = pdf.get_y()
                pdf.line(x, y, x, y + 12)
                pdf.set_x(x + 6)
                pdf.multi_cell(0, 6, _safe_pdf_text(el.get_text(strip=True)))
                pdf.ln(4)

            elif tag == "table":
                self._render_table(pdf, el)

            elif tag == "div":
                # Recurse into divs (BlockNote wraps content in divs)
                self._render_elements(pdf, el)

            else:
                # Fallback: treat as paragraph
                text = el.get_text(strip=True)
                if text:
                    pdf.set_font("Helvetica", "", 11)
                    pdf.set_text_color(51, 51, 51)
                    pdf.multi_cell(0, 6, _safe_pdf_text(text))
                    pdf.ln(3)

    def _render_table(self, pdf: _PremiumPDF, table_el: Tag) -> None:
        """Render an HTML table into the PDF."""
        rows = table_el.find_all("tr")
        if not rows:
            return
        col_count = max(
            (len(r.find_all(["td", "th"])) for r in rows if isinstance(r, Tag)),
            default=0,
        )
        if col_count == 0:
            return
        col_w = (pdf.w - 40) / col_count
        for row in rows:
            if not isinstance(row, Tag):
                continue
            cells = row.find_all(["td", "th"])
            first_cell = cells[0] if cells else None
            is_header = first_cell.name == "th" if isinstance(first_cell, Tag) else False
            for cell in cells:
                if not isinstance(cell, Tag):
                    continue
                if is_header:
                    pdf.set_font("Helvetica", "B", 10)
                    pdf.set_fill_color(240, 240, 240)
                else:
                    pdf.set_font("Helvetica", "", 10)
                    pdf.set_fill_color(255, 255, 255)
                pdf.set_text_color(51, 51, 51)
                text = _safe_pdf_text(cell.get_text(strip=True))[:40]
                pdf.cell(col_w, 7, text, border=1, fill=True)
            pdf.ln()
        pdf.ln(4)

    # ----------------------------------------------------------------- DOCX
    def generate_docx(self, html_content: str, title: str = "DeepSpace Note") -> io.BytesIO:
        """Generates a structured DOCX from HTML content."""
        doc = Document()
        doc.add_heading(title, 0)

        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.find_all(["h1", "h2", "h3", "p", "ul", "ol", "pre"]):
            if not isinstance(element, Tag):
                continue
            if element.name == "h1":
                doc.add_heading(element.get_text(), level=1)
            elif element.name == "h2":
                doc.add_heading(element.get_text(), level=2)
            elif element.name == "h3":
                doc.add_heading(element.get_text(), level=3)
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name in ["ul", "ol"]:
                for li in element.find_all("li"):
                    if not isinstance(li, Tag):
                        continue
                    style = "List Bullet" if element.name == "ul" else "List Number"
                    doc.add_paragraph(li.get_text(), style=style)
            elif element.name == "pre":
                doc.add_paragraph(element.get_text(), style="No Spacing")

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output

    # ------------------------------------------------------------------ MD
    def generate_md(self, html_content: str) -> io.BytesIO:
        """Convert the note HTML into a portable Markdown document.

        BlockNote stores the note as HTML, so returning that HTML with a ``.md``
        extension produces a file that is not actually Markdown.  Keep the
        conversion local and deterministic so exports do not depend on a
        browser or an additional service.
        """
        soup = BeautifulSoup(html_content, "html.parser")
        lines: list[str] = []

        def inline(node: object) -> str:
            if not isinstance(node, Tag):
                return str(node)
            tag = node.name.lower()
            text = "".join(inline(child) for child in node.children)
            if tag in {"strong", "b"}:
                return f"**{text.strip()}**"
            if tag in {"em", "i"}:
                return f"*{text.strip()}*"
            if tag == "code":
                return f"`{text.strip()}`"
            if tag == "a":
                href = str(node.get("href", "")).strip()
                return f"[{text.strip()}]({href})" if href else text
            if tag == "br":
                return "\n"
            return text

        def render(node: object, level: int = 0) -> None:
            if not isinstance(node, Tag):
                return
            tag = node.name.lower()
            if tag in {"h1", "h2", "h3", "h4", "h5", "h6"}:
                lines.append(f"{'#' * int(tag[1])} {inline(node).strip()}\n")
            elif tag == "p":
                value = inline(node).strip()
                if value:
                    lines.append(f"{value}\n")
            elif tag in {"ul", "ol"}:
                ordered = tag == "ol"
                index = 1
                for child in node.find_all("li", recursive=False):
                    marker = f"{index}." if ordered else "-"
                    lines.append(f"{marker} {inline(child).strip()}\n")
                    index += 1
                lines.append("\n")
            elif tag == "blockquote":
                value = inline(node).strip()
                if value:
                    lines.append("\n".join(f"> {line}" for line in value.splitlines()) + "\n")
            elif tag == "pre":
                lines.append(f"```\n{node.get_text().rstrip()}\n```\n")
            elif tag == "table":
                rows = node.find_all("tr")
                matrix = [
                    [inline(cell).strip() for cell in row.find_all(["th", "td"])] for row in rows
                ]
                if matrix:
                    width = max(len(row) for row in matrix)
                    matrix = [row + [""] * (width - len(row)) for row in matrix]
                    lines.append("| " + " | ".join(matrix[0]) + " |\n")
                    lines.append("| " + " | ".join("---" for _ in range(width)) + " |\n")
                    for row in matrix[1:]:
                        lines.append("| " + " | ".join(row) + " |\n")
                    lines.append("\n")
            elif tag in {"div", "section", "article", "main", "body"}:
                for child in node.children:
                    render(child, level)

        for child in soup.children:
            render(child)

        output = io.BytesIO()
        output.write("".join(lines).strip().encode("utf-8"))
        output.seek(0)
        return output
